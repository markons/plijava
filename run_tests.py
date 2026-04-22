"""
Test runner for plijava.py.
Runs each .pli test file through the transpiler and collects results.
Writes a report to c:/temp/plijava_tests.docx.
"""
import subprocess
import sys
import os
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

WRAPPER  = r'c:\githubs\plijava\plijava_wrapper.py'
TEST_DIR = r'c:\githubs\plijava\tests'
PYTHON   = sys.executable
TIMEOUT  = 60   # seconds per test

# ── test catalogue ──────────────────────────────────────────────────────────
TESTS = [
    ('test01_basic.pli',      'Variable declarations, assignments, put skip list'),
    ('test02_arithmetic.pli', 'Arithmetic expressions: + - * /'),
    ('test03_if.pli',         'if-then-else conditional'),
    ('test04_select.pli',     'select-when-other (case/switch)'),
    ('test05_dowhile.pli',    'do-while loop'),
    ('test06_dofromto.pli',   'do-from-to counted loop'),
    ('test07_array1d.pli',    '1-D integer array'),
    ('test08_array2d.pli',    '2-D integer array'),
    ('test09_concat.pli',     'String concatenation via ||'),
    ('test10_substr.pli',     'Built-in: substr(s, start, len)'),
    ('test11_index.pli',      'Built-in: index(s, char)'),
    ('test12_mod.pli',        'Built-in: mod(x, n)'),
    ('test13_decimal.pli',    'Built-in: decimal(x) -> String'),
    ('test14_random.pli',     'Built-in: random(n)'),
    ('test15_intfunc.pli',    'Internal function returning fixed bin'),
    ('test16_charfunc.pli',   'Internal function returning char'),
    ('test17_voidproc.pli',   'Internal void procedure + call statement'),
    ('test18_fileio.pli',     'File I/O: open / write / read / close'),
    ('test19_doblock.pli',    'do;...end; block inside if-then-else'),
    # ---- v2.00 features (Claude Code) ----
    ('test20_leave.pli',         'LEAVE statement (break out of do-while)'),
    ('test21_iterate.pli',       'ITERATE statement (continue in do-while)'),
    ('test22_stop.pli',          'STOP statement -> System.exit(0)'),
    ('test23_doby.pli',          'DO FROM/TO BY (positive and negative step)'),
    ('test24_init.pli',          'DCL ... INIT(value) initializer'),
    ('test25_bit.pli',           'BIT(1) type -> Java boolean'),
    ('test26_strbuiltins.pli',   'String built-ins: length, trim, upper, lower, repeat, reverse'),
    ('test27_mathbuiltins.pli',  'Math built-ins: abs, sqrt, ceil, floor, round, trunc, sign, max, min'),
    ('test28_datetime.pli',      'date() and time() built-ins'),
    ('test29_endfile.pli',       'ON ENDFILE handler with read loop'),
]

# ── helpers ─────────────────────────────────────────────────────────────────
def extract(output, pattern, default=''):
    m = re.search(pattern, output, re.DOTALL)
    return m.group(1).strip() if m else default

def run_test(pli_file):
    path = os.path.join(TEST_DIR, pli_file)
    try:
        proc = subprocess.run(
            [PYTHON, WRAPPER, path],
            capture_output=True, text=True, timeout=TIMEOUT
        )
        out = proc.stdout + proc.stderr
    except subprocess.TimeoutExpired:
        return 'TIMEOUT', '', '', 'Process exceeded timeout'
    except Exception as e:
        return 'ERROR', '', '', str(e)

    # ── parse result ──
    if 'Parsing failed.' in out:
        parse = 'FAIL'
    elif 'Syntax error' in out:
        parse = 'FAIL'
    else:
        parse = 'OK'

    # ── compile result ──
    compile_err = extract(out, r'Compilation failed:(.*?)(?=\n===|\Z)', '')
    if compile_err:
        compile_r = 'FAIL'
    elif 'javac' in out.lower() and 'error' in out.lower():
        compile_r = 'FAIL'
        compile_err = extract(out, r'(error:.*)', '')
    else:
        compile_r = 'OK' if parse == 'OK' else 'N/A'

    # ── execution result ──
    exec_out = extract(out, r'===Execution result:={20,}\s*(.*?)(?:={10,}|$)', '')
    exec_fail = extract(out, r'Execution failed:(.*?)(?:\n\n|\Z)', '')
    if exec_fail:
        exec_out = 'FAIL: ' + exec_fail[:200]
    elif not exec_out and compile_r == 'OK':
        exec_out = '(no output)'

    # ── overall status ──
    if parse == 'FAIL':
        status = 'PARSE FAIL'
    elif compile_r == 'FAIL':
        status = 'COMPILE FAIL'
    elif exec_fail:
        status = 'RUNTIME FAIL'
    elif compile_r == 'OK' and not exec_fail:
        status = 'PASS'
    else:
        status = 'UNKNOWN'

    notes = (compile_err or '').strip()[:300]
    return status, parse, compile_r, exec_out[:300], notes

# ── colour helpers ───────────────────────────────────────────────────────────
STATUS_COLOR = {
    'PASS':         RGBColor(0x00, 0x80, 0x00),   # green
    'COMPILE FAIL': RGBColor(0xC0, 0x00, 0x00),   # red
    'PARSE FAIL':   RGBColor(0xC0, 0x00, 0x00),   # red
    'RUNTIME FAIL': RGBColor(0xFF, 0x80, 0x00),   # orange
    'TIMEOUT':      RGBColor(0xFF, 0x80, 0x00),
    'ERROR':        RGBColor(0xC0, 0x00, 0x00),
    'UNKNOWN':      RGBColor(0x80, 0x80, 0x80),
}

def colored_cell(cell, text, color):
    para = cell.paragraphs[0]
    para.clear()
    run = para.add_run(text)
    run.bold = True
    run.font.color.rgb = color

def small(cell, text):
    para = cell.paragraphs[0]
    para.clear()
    run = para.add_run(text)
    run.font.size = Pt(8)

# ── main ─────────────────────────────────────────────────────────────────────
print('Running tests...')
results = []
for pli_file, description in TESTS:
    print(f'  {pli_file} ...', end='', flush=True)
    result = run_test(pli_file)
    results.append((pli_file, description) + result)
    print(result[0])

# ── build docx ───────────────────────────────────────────────────────────────
doc = Document()

title = doc.add_heading('plijava.py — Feature Test Report', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

from datetime import date
doc.add_paragraph(f'Generated: {date.today()}   |   Python: {sys.version.split()[0]}   |   Tests: {len(TESTS)}')

# summary counts
passed  = sum(1 for r in results if r[2] == 'PASS')
pfail   = sum(1 for r in results if r[2] == 'PARSE FAIL')
cfail   = sum(1 for r in results if r[2] == 'COMPILE FAIL')
rfail   = sum(1 for r in results if r[2] == 'RUNTIME FAIL')
other   = len(results) - passed - pfail - cfail - rfail

doc.add_paragraph(
    f'PASS: {passed}   PARSE FAIL: {pfail}   COMPILE FAIL: {cfail}   '
    f'RUNTIME FAIL: {rfail}   OTHER: {other}'
)

doc.add_heading('Detailed Results', level=1)

table = doc.add_table(rows=1, cols=6)
table.style = 'Table Grid'
headers = ['Test file', 'Feature', 'Status', 'Parse', 'Compile', 'Java output / Error']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.paragraphs[0].clear()
    run = cell.paragraphs[0].add_run(h)
    run.bold = True

for pli_file, description, status, parse, compile_r, exec_out, notes in results:
    row = table.add_row().cells
    row[0].text = pli_file
    row[1].text = description
    colored_cell(row[2], status, STATUS_COLOR.get(status, RGBColor(0,0,0)))
    row[3].text = parse
    row[4].text = compile_r
    detail = exec_out if exec_out else notes
    small(row[5], detail[:300] if detail else '')

# ── set column widths ─────────────────────────────────────────────────────────
from docx.shared import Inches
widths = [1.3, 2.0, 1.0, 0.6, 0.7, 2.8]
for row in table.rows:
    for i, cell in enumerate(row.cells):
        cell.width = Inches(widths[i])

# ── notes section ─────────────────────────────────────────────────────────────
doc.add_heading('Known Bugs Observed', level=1)
bugs = [
    ('PARSE FAIL tests',
     'Grammar reduce/reduce conflicts (declaration_list left+right recursive, '
     'expression:ID duplicates variable_access:ID) can cause unexpected parse failures.'),
    ('COMPILE FAIL — substr',
     'SUBSTR with single start argument generates .substring(n:) — Python slice '
     'syntax, not valid Java. Fix: change to .substring(n).'),
    ('COMPILE FAIL — SQL',
     'exec sql assignment generates double semicolon: var = Integer.parseInt(result);;'),
    ('t_FILENAME / VARYING',
     'FILENAME not in tokens tuple; VARYING in reserved dict but not in tokens. '
     'PLY may emit a LexError or warning at startup.'),
    ('put skip list output',
     'Elements joined with "+ " only — no spaces between numeric values in the '
     'Java println output.'),
]
for title_b, body in bugs:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(title_b + ': ')
    run.bold = True
    p.add_run(body)

os.makedirs('c:/temp', exist_ok=True)
doc.save('c:/temp/plijava_tests.docx')
print('\nReport written to c:/temp/plijava_tests.docx')
